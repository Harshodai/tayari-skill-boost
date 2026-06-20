"""Builds an ATS-friendly .docx file from optimized resume text.
Understands a light markdown-ish structure: headings, bullets, plain paragraphs.
"""
import io
import re

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def _clean_inline(text: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text).strip()


def build_resume_docx(resume_text: str, title: str = "Resume") -> io.BytesIO:
    if not HAS_DOCX:
        # Fallback: return plain text as bytes with a .docx mime type
        buf = io.BytesIO()
        buf.write(resume_text.encode("utf-8"))
        buf.seek(0)
        return buf

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    lines = resume_text.splitlines()
    first_content_line = True
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        stripped = line.strip()

        is_md_heading = stripped.startswith("#")
        is_caps_heading = (
            len(stripped) < 60 and stripped.upper() == stripped
            and re.search("[A-Z]{3,}", stripped) and not stripped.startswith(("-", "•", "*"))
        )
        if is_md_heading or is_caps_heading:
            text = _clean_inline(stripped.lstrip("#"))
            if first_content_line:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(16)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_paragraph()
                run = p.add_run(text.upper())
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
                p.paragraph_format.space_before = Pt(8)
            first_content_line = False
            continue

        if stripped.startswith(("- ", "• ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(_clean_inline(stripped[2:]))
        else:
            p = doc.add_paragraph()
            run = p.add_run(_clean_inline(stripped))
            if first_content_line:
                run.bold = True
                run.font.size = Pt(15)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        first_content_line = False

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
